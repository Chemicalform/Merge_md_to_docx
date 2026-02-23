# -*- coding: utf-8 -*-
"""
遍历指定根目录及所有子目录，对每个包含 .md 的文件夹单独生成一个 docx，
整合该文件夹下的所有 .md 文件，按文件名分章节，正文两端对齐，章节间分页。

深度性能优化：
1. 内存中将所有 Markdown 拼接为文本流，一次性无缝传递给 Pandoc。
2. 通过 Raw OpenXML 注入原生 Word 分页符，取代缓慢的 Python 层面深拷贝拼接。
3. 全局缓存正则表达式与 XML 命名空间，实现 XPath 零开销就地 (In-place) 清洗。

依赖: pip install python-docx pypandoc_binary
"""
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from docx.oxml.ns import qn
from concurrent.futures import ProcessPoolExecutor, as_completed
import pypandoc
import tempfile
import os
import re
import argparse
import shutil

# ---------------------------------------------------------------------------
# 全局缓存 (性能优化极速清洗所需)
# ---------------------------------------------------------------------------
_SPACE_RE = re.compile(r"[ \t]{2,}")
_W_BR = qn("w:br")
_W_TYPE = qn("w:type")

# Pandoc 专用的原生 OpenXML 分页符注入片段
_PAGE_BREAK_MD = '\n\n```{=openxml}\n<w:p><w:r><w:br w:type="page"/></w:r></w:p>\n```\n\n'

# ---------------------------------------------------------------------------
# 工具函数
# ---------------------------------------------------------------------------

def get_md_files(folder: Path, ext=".md"):
    """获取目录下的指定后缀文件（默认 .md），按文件名排序"""
    if not ext.startswith("."): ext = "." + ext
    files = [f for f in folder.iterdir() if f.is_file() and f.suffix.lower() == ext.lower()]
    return sorted(files, key=lambda p: p.name.lower())


def read_text_with_fallback(path, primary="utf-8", fallbacks=("gbk", "gb2312", "latin-1")):
    """尝试多种编码读取文本文件，优先 UTF-8，失败后依次回退。"""
    try:
        return path.read_text(encoding=primary)
    except UnicodeDecodeError:
        for enc in fallbacks:
            try:
                return path.read_text(encoding=enc)
            except UnicodeDecodeError:
                continue
    return path.read_text(encoding=primary, errors="replace")


# ---------------------------------------------------------------------------
# DOCX 就地后处理 (In-place Mutation 极速替换版)
# ---------------------------------------------------------------------------

def _optimize_paragraph(para, justify: bool):
    """就地优化单个段落：可选两端对齐、拔除软换行、压缩多余空格"""
    if justify:
        try:
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        except Exception:
            pass

    p_element = para._element
    # 直接深度查找能够跳过 python-docx runs 迭代器的盲点（如超链接内嵌套的 w:br）
    for br in p_element.findall('.//' + _W_BR):
        br_type = br.get(_W_TYPE)
        if br_type is None or br_type == "textWrapping":
            br.getparent().remove(br)

    for t in p_element.findall('.//' + qn('w:t')):
        if t.text:
            t.text = _SPACE_RE.sub(" ", t.text)


def _set_style_font(style, font_name, pt_size):
    """设置样式的中英文字体与字号"""
    if pt_size:
        style.font.size = Pt(pt_size)
    if font_name:
        style.font.name = font_name
        # python-docx 中设置 font.name 会生成 rPr 和 w:rFonts 节点
        # 这里补充设置 w:eastAsia 以支持中文字体
        rPr = style.font._element.get_or_add_rPr()
        if rPr.rFonts is not None:
            rPr.rFonts.set(qn('w:eastAsia'), font_name)


def _is_empty_paragraph(para):
    """判断段落是否为真正的空行（无文字、无分页、无图片）"""
    if para.text.strip():
        return False
    el = para._element
    if el.findall('.//' + _W_BR): return False
    if el.findall('.//' + qn('w:drawing')): return False
    if el.findall('.//' + qn('w:pict')): return False
    return True


def post_process_docx(doc_path, margin_cm, font_name="Aptos", font_size=12):
    """加载 Pandoc 生成的成型文档，只进行一层 O(N) 遍历做就地清扫，随后直接保存。"""
    doc = Document(doc_path)
    margin = Cm(margin_cm)

    # 0. 设置全局字体样式 (O(1) 极速修改)
    try:
        # 遍历覆盖所有的文档样式，强制把字体替换为要求的全局字体，包含 Pandoc 用于代码块的 Verbatim Char 和 Source Code
        for s in doc.styles:
            try:
                if s.name and not s.name.startswith('Heading'):
                    _set_style_font(s, font_name, None)
            except Exception:
                pass

        # 正文基本样式：额外锁定字号
        for s_name in ['Normal', 'Body Text', 'First Paragraph', 'List Paragraph']:
            if s_name in doc.styles:
                _set_style_font(doc.styles[s_name], font_name, font_size)

        # 标题样式 (默认比正文大，这里可根据需求统一配置，为保留层级这里仅将一级标题设为 16 号)
        if 'Heading 1' in doc.styles:
            _set_style_font(doc.styles['Heading 1'], font_name, 16)
        if 'Heading 2' in doc.styles:
            _set_style_font(doc.styles['Heading 2'], font_name, 15)
        if 'Heading 3' in doc.styles:
            _set_style_font(doc.styles['Heading 3'], font_name, 14)
    except Exception as e:
        print(f"Warning: 设置样式字体时出错 {e}")

    # 1. 统一调整页边距
    for section in doc.sections:
        section.top_margin = section.bottom_margin = margin
        section.left_margin = section.right_margin = margin

    # 2. 清扫正文 (含对齐、剔除软换行、剔除空行)
    for para in doc.paragraphs:
        _optimize_paragraph(para, justify=True)
        # 如果处理后段落里既没有文字，也没有任何图片或分页符标记，则当做多余的纯空行删掉
        if _is_empty_paragraph(para):
            para._element.getparent().remove(para._element)

    # 3. 清扫表格 (可选对齐，部分需求中表格内可能只需要清空软换行)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _optimize_paragraph(para, justify=True)
                    if _is_empty_paragraph(para):
                        para._element.getparent().remove(para._element)

    doc.save(doc_path)


# ---------------------------------------------------------------------------
# 核心：处理单个文件夹拼接流
# ---------------------------------------------------------------------------

def process_folder(folder, root, margin_cm, font_name, font_size, ext=".md"):
    """单线程/进程：收集文件夹的文件 -> 组装成一个巨大字符串 -> Pandoc 图转 -> 后处理"""
    md_files = get_md_files(folder, ext)
    if not md_files:
        return None

    # 1. 在内存中融合全部章节
    combined_md_parts = []
    for idx, md_path in enumerate(md_files):
        if idx > 0:
            combined_md_parts.append(_PAGE_BREAK_MD)

        chapter_name = md_path.stem
        combined_md_parts.append(f"# {chapter_name}\n\n")
        combined_md_parts.append(read_text_with_fallback(md_path))
        combined_md_parts.append("\n")

    combined_md = "".join(combined_md_parts)

    # 2. 持久化存储
    fd, tmp_path = tempfile.mkstemp(suffix=".docx")
    os.close(fd)

    try:
        # 3. 调用底层的 Pandoc 进制程序进行 O(1) 转换
        pypandoc.convert_text(
            combined_md,
            "docx",
            format="markdown",
            outputfile=tmp_path
        )

        # 4. 后处理：XML 原位清洗（软换行剥离、边距赋予、对齐规整、全局字体）
        post_process_docx(tmp_path, margin_cm, font_name, font_size)

        # 5. 转移落库
        out_path = folder / f"{folder.name}.docx"
        shutil.move(tmp_path, out_path)

        rel = out_path.relative_to(root)
        return (rel, len(md_files))

    finally:
        # 安全清理遗留产生的 tmp
        try:
            if os.path.exists(tmp_path):
                os.unlink(tmp_path)
        except OSError:
            pass


# ---------------------------------------------------------------------------
# 主入口
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(description="合并目录下的 Markdown 文件为 DOCX 文档 (极致性能版)")
    parser.add_argument("--dir", type=Path, default=None,
                        help="根目录（默认为脚本所在目录）")
    parser.add_argument("--margin", type=float, default=1.27,
                        help="页边距，单位 cm（默认 1.27）")
    parser.add_argument("--encoding", default="utf-8",
                        help="主要编码（默认 utf-8，失败自动回退 GBK 等）")
    parser.add_argument("--workers", type=int, default=None,
                        help="并发进程数（默认为 CPU 核心数）")
    parser.add_argument("--font", type=str, default="Aptos",
                        help="全局默认字体（默认: Aptos）")
    parser.add_argument("--font-size", type=float, default=12,
                        help="全局正文字号（默认 12 号字体）")
    parser.add_argument("--ext", type=str, default=".md",
                        help="要合并的文件后缀名（默认: .md，可改为 .txt 等）")
    args = parser.parse_args()

    root = args.dir if args.dir else Path(__file__).resolve().parent

    # 收集所有包含指定后缀文件的目录
    dirs_with_files = set()
    for root_dir, _, _ in os.walk(root):
        folder = Path(root_dir)
        if get_md_files(folder, args.ext):
            dirs_with_files.add(folder)
    dirs_with_files = sorted(list(dirs_with_files), key=lambda p: (len(p.parts), str(p)))

    if not dirs_with_files:
        print(f"未在脚本目录及子目录下发现任何 {args.ext} 文件。")
        return

    total = len(dirs_with_files)
    workers = args.workers or min(total, os.cpu_count() or 4)
    print(f"发现 {total} 个文件夹，使用 {workers} 个进程并发极速转换…\n")

    completed = 0
    with ProcessPoolExecutor(max_workers=workers) as executor:
        future_to_folder = {
            executor.submit(process_folder, folder, root, args.margin, args.font, args.font_size, args.ext): folder
            for folder in dirs_with_files
        }
        for future in as_completed(future_to_folder):
            folder = future_to_folder[future]
            completed += 1
            try:
                result = future.result()
                if result:
                    rel, count = result
                    print(f"[{completed}/{total}] ✓ 已生成: {rel}（{count} 个章节）")
                else:
                    print(f"[{completed}/{total}] - 跳过: {folder.name}（无 .md 文件）")
            except Exception as e:
                print(f"[{completed}/{total}] ✗ 失败: {folder.name} — {e}")

    print(f"\n完成！共处理 {total} 个文件夹。")


if __name__ == "__main__":
    main()
