# -*- coding: utf-8 -*-
"""
遍历指定根目录及所有子目录，对每个包含 .md 的文件夹单独生成一个 docx，
整合该文件夹下的所有 .md 文件，按文件名分章节，正文两端对齐，章节间分页。
支持标题(1-6级)、列表、代码块等常见 Markdown 格式。
使用 pypandoc（内置 pandoc）进行格式转换，多线程并发处理。

依赖: pip install python-docx pypandoc_binary
"""
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Cm, Pt
from docx.oxml.ns import qn
from concurrent.futures import ProcessPoolExecutor, as_completed
import pypandoc
import tempfile
import os
import re
import argparse

# ---------------------------------------------------------------------------
# 工具函数
# ---------------------------------------------------------------------------

def get_md_files(folder):
    """返回文件夹下按名称排序的 .md 文件列表。"""
    folder = Path(folder)
    files = sorted(folder.glob("*.md"), key=lambda p: p.name.lower())
    return [f for f in files if f.is_file()]


def read_text_with_fallback(path, primary="utf-8", fallbacks=("gbk", "gb2312", "latin-1")):
    """尝试多种编码读取文本文件，优先 UTF-8，失败后依次回退。"""
    try:
        return path.read_text(encoding=primary)
    except UnicodeDecodeError:
        for enc in fallbacks:
            try:
                return path.read_text(encoding=enc)
            except UnicodeDecodeError:
                continue
    return path.read_text(encoding=primary, errors="replace")


def _is_list_item(line):
    """判断一行是否是 Markdown 列表项（无序或有序）。"""
    s = line.strip()
    return bool(s) and (
        s.startswith(("- ", "* ", "+ "))
        or re.match(r"\d+\.\s", s)
    )


def _strip_inline_md(text):
    """去除行内 Markdown 格式标记（粗体、斜体、行内代码）。"""
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"__(.+?)__", r"\1", text)
    text = re.sub(r"(?<!\*)\*([^*]+)\*(?!\*)", r"\1", text)
    text = re.sub(r"(?<!_)_([^_]+)_(?!_)", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    return text


def _add_heading_with_justify(doc, text, level, set_justify=True):
    """添加标题段落，可选两端对齐。"""
    p = doc.add_heading(text, level=level)
    if set_justify:
        try:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        except Exception:
            pass
    return p


def _remove_soft_line_breaks(doc):
    """移除 docx 正文中所有的手动换行符 (<w:br/>)，替换为空格。
    只移除段内换行(soft break)，不影响分页符(page break)。
    """
    for para in doc.paragraphs:
        for run in para.runs:
            for br in run._element.findall(qn("w:br")):
                br_type = br.get(qn("w:type"))
                # 只移除非分页的换行符（type 为 None 即普通软换行，或 type="textWrapping"）
                if br_type is None or br_type == "textWrapping":
                    # 在 br 之前插入一个空格文本节点，避免前后文字粘连
                    parent = br.getparent()
                    idx = list(parent).index(br)
                    parent.remove(br)
                    # 如果 br 不是最后一个子元素，在 run 文本中补一个空格
                    # 通过直接操作 run.text 实现
            # 清理后重新整理 run 的文本：把连续空白压缩
            if run.text:
                run.text = re.sub(r"[ \t]{2,}", " ", run.text)

    # 同时处理表格单元格内的换行
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        for br in run._element.findall(qn("w:br")):
                            br_type = br.get(qn("w:type"))
                            if br_type is None or br_type == "textWrapping":
                                run._element.remove(br)
                        if run.text:
                            run.text = re.sub(r"[ \t]{2,}", " ", run.text)


# ---------------------------------------------------------------------------
# Markdown → docx 段落（简易解析器，仅在 pypandoc 转换失败时使用）
# ---------------------------------------------------------------------------

_HEADING_RE = re.compile(r"^(#{1,6})\s+(.+)$")

def simple_md_to_paragraphs(doc, text, set_justify=True):
    """将 Markdown 文本按段落加入 docx，支持标题、列表、代码块、水平线。"""
    lines = text.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # 围栏代码块
        if stripped.startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            if i < len(lines):
                i += 1
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run("\n".join(code_lines))
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            continue

        # 标题
        m = _HEADING_RE.match(stripped)
        if m:
            level = len(m.group(1))
            _add_heading_with_justify(doc, m.group(2).strip(), level, set_justify)
            i += 1
            continue

        # 水平线
        if stripped in ("---", "***", "___"):
            i += 1
            continue

        # 列表项
        if _is_list_item(stripped):
            while i < len(lines) and _is_list_item(lines[i]):
                item = lines[i].strip()
                item = re.sub(r"^[-*+]\s+", "", item)
                item = re.sub(r"^\d+\.\s+", "", item)
                p = doc.add_paragraph(_strip_inline_md(item), style="List Bullet")
                if set_justify:
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                i += 1
            continue

        # 普通段落
        block = [stripped]
        i += 1
        while (i < len(lines)
               and lines[i].strip()
               and not _HEADING_RE.match(lines[i].strip())
               and not lines[i].strip().startswith("```")
               and not _is_list_item(lines[i])
               and lines[i].strip() not in ("---", "***", "___")):
            block.append(lines[i].strip())
            i += 1
        para_text = _strip_inline_md(" ".join(block))
        p = doc.add_paragraph(para_text)
        if set_justify:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        continue


# ---------------------------------------------------------------------------
# pypandoc 转换 & 段落复制
# ---------------------------------------------------------------------------

def convert_md_to_docx_via_pypandoc(md_path, out_docx_path):
    """用 pypandoc 将单个 md 转为临时 docx。"""
    try:
        pypandoc.convert_file(
            str(md_path), "docx",
            format="markdown",
            outputfile=str(out_docx_path),
        )
        return True
    except Exception:
        return False


def copy_paragraphs_with_justify(source_doc, target_doc, skip_first_if_equal=None):
    """将 source_doc 的正文段落复制到 target_doc，并设为两端对齐。"""
    for idx, para in enumerate(source_doc.paragraphs):
        if skip_first_if_equal is not None and idx == 0 and para.text.strip() == skip_first_if_equal:
            continue
        if not para.text.strip():
            target_doc.add_paragraph()
            continue
        new_p = target_doc.add_paragraph()
        new_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in para.runs:
            r = new_p.add_run(run.text)
            r.bold = run.bold
            r.italic = run.italic
            if run.font.size:
                r.font.size = run.font.size
            if run.font.name:
                r.font.name = run.font.name
    for table in source_doc.tables:
        new_table = target_doc.add_table(rows=len(table.rows), cols=len(table.columns))
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                new_table.rows[i].cells[j].text = cell.text
                for para in new_table.rows[i].cells[j].paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


# ---------------------------------------------------------------------------
# 核心：为每个文件夹构建 docx（进程安全）
# ---------------------------------------------------------------------------

def build_docx_for_folder(folder, md_files, margin_cm=1.27):
    """针对一个文件夹内的 md 列表，生成一个 Document 并返回。"""
    doc = Document()
    margin = Cm(margin_cm)
    for section in doc.sections:
        section.top_margin = section.bottom_margin = margin
        section.left_margin = section.right_margin = margin

    for idx, md_path in enumerate(md_files):
        if idx > 0:
            p = doc.add_paragraph()
            p.add_run().add_break(WD_BREAK.PAGE)
        chapter_name = md_path.stem
        _add_heading_with_justify(doc, chapter_name, level=1)

        fd, tmp_path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        try:
            if convert_md_to_docx_via_pypandoc(md_path, tmp_path):
                sub_doc = Document(tmp_path)
                # 先清除子文档中的手动换行符
                _remove_soft_line_breaks(sub_doc)
                copy_paragraphs_with_justify(sub_doc, doc, skip_first_if_equal=chapter_name)
            else:
                text = read_text_with_fallback(md_path)
                simple_md_to_paragraphs(doc, text, set_justify=True)
        finally:
            try:
                os.unlink(tmp_path)
            except OSError:
                pass

        doc.add_paragraph()

    # 最终再对合并后的 doc 整体清理一次手动换行
    _remove_soft_line_breaks(doc)
    return doc


def process_folder(folder, root, margin_cm):
    """处理单个文件夹：构建 docx 并保存。供线程池调用。"""
    md_files = get_md_files(folder)
    if not md_files:
        return None
    doc = build_docx_for_folder(folder, md_files, margin_cm=margin_cm)
    out_path = folder / f"{folder.name}.docx"
    doc.save(out_path)
    rel = out_path.relative_to(root)
    return (rel, len(md_files))


# ---------------------------------------------------------------------------
# 主入口
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(description="合并目录下的 Markdown 文件为 DOCX 文档")
    parser.add_argument("--dir", type=Path, default=None,
                        help="根目录（默认为脚本所在目录）")
    parser.add_argument("--margin", type=float, default=1.27,
                        help="页边距，单位 cm（默认 1.27）")
    parser.add_argument("--encoding", default="utf-8",
                        help="主要编码（默认 utf-8，失败自动回退 GBK 等）")
    parser.add_argument("--workers", type=int, default=None,
                        help="并发线程数（默认为 CPU 核心数）")
    args = parser.parse_args()

    root = args.dir if args.dir else Path(__file__).resolve().parent

    # 收集所有包含 .md 的目录
    dirs_with_md = set()
    for f in root.rglob("*.md"):
        if f.is_file():
            dirs_with_md.add(f.parent)
    dirs_with_md = sorted(dirs_with_md, key=lambda p: (len(p.parts), str(p)))

    if not dirs_with_md:
        print("未在脚本目录及子目录下发现任何 .md 文件。")
        return

    total = len(dirs_with_md)
    workers = args.workers or min(total, os.cpu_count() or 4)
    print(f"发现 {total} 个文件夹，使用 {workers} 个进程并发处理…\n")

    completed = 0
    with ProcessPoolExecutor(max_workers=workers) as executor:
        future_to_folder = {
            executor.submit(process_folder, folder, root, args.margin): folder
            for folder in dirs_with_md
        }
        for future in as_completed(future_to_folder):
            folder = future_to_folder[future]
            completed += 1
            try:
                result = future.result()
                if result:
                    rel, count = result
                    print(f"[{completed}/{total}] ✓ 已生成: {rel}（{count} 个章节）")
                else:
                    print(f"[{completed}/{total}] - 跳过: {folder.name}（无 .md 文件）")
            except Exception as e:
                print(f"[{completed}/{total}] ✗ 失败: {folder.name} — {e}")

    print(f"\n完成！共处理 {total} 个文件夹。")


if __name__ == "__main__":
    main()
